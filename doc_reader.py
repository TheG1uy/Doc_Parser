from docx import Document

class doc_reader:
    def __init__(self, _path_to_doc):
        self.doc = None
        if (_path_to_doc != None):
            self.path_to_doc = _path_to_doc
            self.doc = Document(self.path_to_doc)

    def print_doc(self):
        results = []
        for line_num in range(len(self.doc.paragraphs)):
                results.append(self.doc.paragraphs[line_num].text)
        return results

    def get_doc(self):
        return self.doc

    def save_file(self):
        self.doc.save(self.path_to_doc)
        return 'Документ ' + self.path_to_doc +  ' сохранен'